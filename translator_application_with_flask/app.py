#If we execute app.py in our terminal, the code will run in local host, by taking that local host ip address we can test the code in any browser.
#We need to install the required libraries by using the below command
#pip install Flask==2.0.2 googletrans==4.0.0-rc1 python-docx==0.8.11 PyPDF2==1.26.0

# Importing all the required libraries
import os
from flask import Flask, request, jsonify, send_file, render_template
from googletrans import Translator
from docx import Document
from PyPDF2 import PdfFileReader
from io import BytesIO

app = Flask(__name__)

# Defining route for home page
@app.route('/')
def home():
    return render_template('index.html')

# Defining route for translation
@app.route('/translate', methods=['POST'])
def translate():
    if 'file' in request.files:
        file = request.files['file']
        input_language = request.form['input-language']
        translated_language = request.form['translated-language']
        if file.filename.endswith('.pdf'):
            translated_text = translate_pdf(file, input_language, translated_language)
            return save_pdf(translated_text)
        elif file.filename.endswith('.docx'):
            translated_text = translate_docx(file, input_language, translated_language)
            return save_docx(translated_text)

    elif 'text' in request.form:
        text = request.form['text']
        translated_language = request.form['translated-language']
        translated_text = translate_text(text, translated_language)
        return render_template('translated.html', translated_text=translated_text)

    return jsonify({'error': 'Unsupported input format'})

def translate_pdf(file, input_language, translated_language):
    translator = Translator()
    translated_text = ""

    pdf_reader = PdfFileReader(file)
    for page_num in range(pdf_reader.numPages):
        page = pdf_reader.getPage(page_num)
        text = page.extractText()
        translated_text += translator.translate(text, src=input_language, dest=translated_language).text

    return translated_text

def translate_docx(file, input_language, translated_language):
    translator = Translator()
    translated_text = ""

    document = Document(file)
    for paragraph in document.paragraphs:
        translated_text += translator.translate(paragraph.text, src=input_language, dest=translated_language).text + "\n"

    return translated_text

def translate_text(text, translated_language):
    translator = Translator()
    translated_text = translator.translate(text, dest=translated_language).text
    return translated_text

def save_pdf(translated_text):
    pdf_bytes = BytesIO()
    pdf_bytes.write(translated_text.encode())
    pdf_bytes.seek(0)
    return send_file(pdf_bytes, attachment_filename='translated_document.pdf', as_attachment=True)

def save_docx(translated_text):
    docx_bytes = BytesIO()
    doc = Document()
    doc.add_paragraph(translated_text)
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return send_file(docx_bytes, attachment_filename='translated_document.docx', as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
