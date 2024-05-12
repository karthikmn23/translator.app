#If we execute translation_app.py in our terminal, the code will run in localhost, by taking that localhost ip address we can test the code in any browser.
# We need to execute  streamlit run translation_app.py in our terminal to execute the below code
# We need to install the required libraries b y using below command
# pip install streamlit googletrans==4.0.0-rc1 python-docx

import streamlit as st
from googletrans import Translator
from docx import Document
from PyPDF2 import PdfFileReader
from io import BytesIO

def translate_pdf(file, input_language, translated_language):
    translator = Translator()
    translated_text = ""

    pdf_reader = PdfFileReader(file)
    for page_num in range(pdf_reader.numPages):
        page = pdf_reader.getPage(page_num)
        text = page.extractText()
        translated_text += translator.translate(text, src=input_language, dest=translated_language).text

    return translated_text

def translate_docx(file, input_language, translated_language):
    translator = Translator()
    translated_text = ""

    document = Document(file)
    for paragraph in document.paragraphs:
        translated_text += translator.translate(paragraph.text, src=input_language, dest=translated_language).text + "\n"

    return translated_text

def translate_text(text, translated_language):
    translator = Translator()
    translated_text = translator.translate(text, dest=translated_language).text
    return translated_text

def save_pdf(translated_text):
    pdf_bytes = BytesIO()
    pdf_bytes.write(translated_text.encode())
    pdf_bytes.seek(0)
    return pdf_bytes

def save_docx(translated_text):
    docx_bytes = BytesIO()
    doc = Document()
    doc.add_paragraph(translated_text)
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes

st.title("Translation App")

translation_type = st.radio("Select input type:", ["Text", "File"])
if translation_type == "Text":
    text_input = st.text_area("Enter Text:")
    translated_language = st.selectbox("Select Translated Language:", ["German", "Italian"])
    translated_language_code = "de" if translated_language == "German" else "it"

    if st.button("Translate"):
        translated_text = translate_text(text_input, translated_language_code)
        st.write("Translated Text:")
        st.write(translated_text)

elif translation_type == "File":
    uploaded_file = st.file_uploader("Upload File:", type=["pdf", "docx"])
    input_language = "en"  # Input language fixed to English
    translated_language = st.selectbox("Select Translated Language:", ["German", "Italian"])
    translated_language_code = "de" if translated_language == "German" else "it"

    if uploaded_file is not None:
        if uploaded_file.type == "application/pdf":
            if st.button("Translate"):
                translated_text = translate_pdf(uploaded_file, input_language, translated_language_code)
                st.write("Translated PDF:")
                st.download_button("Download Translated PDF", save_pdf(translated_text), file_name="translated_document.pdf")

        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            if st.button("Translate"):
                translated_text = translate_docx(uploaded_file, input_language, translated_language_code)
                st.write("Translated DOCX:")
                st.download_button("Download Translated DOCX", save_docx(translated_text), file_name="translated_document.docx")

        else:
            st.write("Unsupported file format. Please upload a PDF or DOCX file.")
